from docx import Document

def update_word(serial, health, os_info, hub, screenshot):
    doc = Document("Protocol.docx")

    for table in doc.tables:
        for row in table.rows:
            if "CT900F" in row.cells[0].text:
                row.cells[1].text = (
                    f"Model: {os_info['model']}\n"
                    f"Android: {os_info['android']}\n"
                    f"Patch: {os_info['patch']}\n"
                    # f"UAT: {hub['uat']}\n"
                    f"UAT Server: {hub['uat']}\n"
                    f"Group ID: {hub['group']}\n"
                    f"Username: {hub['username']}\n"
                    f"Group: {hub['group']}"
                )
                row.cells[1].add_paragraph().add_run().add_picture(screenshot)

    doc.save("Protocol.docx")
